from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/chengfeng/Downloads/未命名文件夹/byte-resume-output/盛俊鹏-前端工程师-字节投递版.docx"


def set_cell_text(cell, text, bold=False, size=9.5, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="E5E7EB", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_margins(section):
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)


def style_run(run, size=9.5, bold=False, color="111827"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", size=9.5, bold=False, color="111827", after=2.2, before=0, line=1.05):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = line
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, after=1.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.03
    run = p.add_run(text)
    style_run(run, size=8.8, color="111827")
    return p


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    style_run(run, size=10.8, bold=True, color="0F172A")
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "CBD5E1")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_role(doc, company, role, dates, desc=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1.5)
    r1 = p.add_run(company)
    style_run(r1, size=9.7, bold=True, color="111827")
    r2 = p.add_run(" | {} | {}".format(role, dates))
    style_run(r2, size=9.1, color="374151")
    if desc:
        add_para(doc, desc, size=8.8, color="374151", after=1.5)


def add_project(doc, name, stack, summary, bullets, result=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(name)
    style_run(r1, size=9.5, bold=True, color="111827")
    r2 = p.add_run(" | {}".format(stack))
    style_run(r2, size=8.8, color="4B5563")
    add_para(doc, summary, size=8.8, color="374151", after=1.2)
    for item in bullets:
        add_bullet(doc, item, after=1.0)
    if result:
        add_bullet(doc, "结果：{}".format(result), after=1.5)


def add_topic(doc, index, title, difficulty, actions, result):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0.1)
    run = p.add_run("{}.{}".format(index, title))
    style_run(run, size=8.4, bold=True, color="111827")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.95
    r = p.add_run("难点:")
    style_run(r, size=7.8, bold=True, color="111827")
    r = p.add_run(difficulty)
    style_run(r, size=7.8, color="111827")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("行动:")
    style_run(r, size=7.8, bold=True, color="111827")

    for action_index, action in enumerate(actions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 0.95
        r = p.add_run("{}){}".format(action_index, action))
        style_run(r, size=7.8, color="111827")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0.4)
    p.paragraph_format.line_spacing = 0.95
    r = p.add_run("结果:")
    style_run(r, size=7.8, bold=True, color="111827")
    r = p.add_run(result)
    style_run(r, size=7.8, color="111827")


doc = Document()
for section in doc.sections:
    set_margins(section)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(9)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(1)
r = title.add_run("盛俊鹏")
style_run(r, size=18, bold=True, color="0F172A")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(4)
r = subtitle.add_run("7年前端工程师 / 软件工程本科 / 画布与设计工程化方向 / AI产品前端")
style_run(r, size=10, color="334155")

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(4)
r = contact.add_run("18186473851 | 1300749625@qq.com | 武汉 / 可沟通北京、杭州、上海 | GitHub: https://github.com/rideWind97")
style_run(r, size=8.6, color="475569")

add_section(doc, "个人优势")
for item in [
    "7年前端经验，近两段在 AI 创作工具和协同设计工具中承担核心画布负责人角色，长期处理大文件渲染、复杂交互、流式反馈和线上稳定性问题。",
    "能从业务目标拆到技术方案并推进落地：主导画布节点系统、连接线拓扑、Canvas Player、D2C重构、性能监控等模块，形成可复用的插件、文档和工程规范。",
    "具备团队级提效经验：在多项目规范分散、复用成本高的背景下，推动 Monorepo、微前端、组件库、Design Token、CI/CD、AICR 落地，带来开发效率和协作质量提升。",
]:
    add_bullet(doc, item)

add_section(doc, "技术关键词")
table = doc.add_table(rows=4, cols=2)
table.autofit = False
labels = [
    ("前端基础", "JavaScript ES6+ / TypeScript / Vue2、Vue3 / React Hooks / Pinia / Vite / Webpack"),
    ("图形画布", "Canvas / LeaferJS / Wasm / 复杂坐标系 / 节点拓扑 / 连接线 / Canvas Player / FPS与内存监控"),
    ("工程架构", "pnpm workspace / Turbo / Monorepo / qiankun / wujie / 组件库 / Storybook / Design Token / Atomic CSS"),
    ("AI与效率", "AI Agent交互 / SSE / WebSocket / Prompt设计 / GitLab Webhook / AICR / CI/CD / Docker / Nginx"),
]
for row, (k, v) in zip(table.rows, labels):
    row.cells[0].width = Inches(1.1)
    row.cells[1].width = Inches(5.9)
    set_cell_text(row.cells[0], k, bold=True, size=8.5, color="0F172A")
    set_cell_text(row.cells[1], v, size=8.5, color="111827")
    set_cell_shading(row.cells[0], "F8FAFC")
    for cell in row.cells:
        set_cell_border(cell)

add_section(doc, "工作经历")
add_role(
    doc,
    "ZMO.AI（灵跃苍穹科技有限公司）",
    "北京项目组前端负责人 / 画布负责人",
    "2025.10 - 至今",
    "负责 Buzzy AI视频生成平台核心画布区、Agent 对话框和前端团队招聘培训。"
)
for item in [
    "在 AI 视频生成流程交互链路长、用户等待感强的背景下，我主导画布节点、动态连接线、Canvas Player 和翻译链路的方案设计与核心编码，让生成过程从黑盒等待变成可视化节点流转。",
    "针对大模型响应不确定、前端状态同步复杂的问题，我设计 SSE / WebSocket 流式反馈与 UI 工具调用联动，使用户输入、节点状态和画布反馈能够实时闭环。",
    "为了降低画布问题排查成本，我制定 FPS、内存泄漏、渲染延迟等监控指标并沉淀模块文档，帮助团队更快定位性能退化和协作边界问题。",
]:
    add_bullet(doc, item)

add_role(
    doc,
    "蓝湖（武汉界面进步有限公司）",
    "画布负责人 / 设计工程化负责人",
    "2024.04 - 2025.10",
    "负责 MasterGo 画布区重点功能、D2C模块重构、官网架构升级、性能治理及组内技术指导。"
)
for item in [
    "在 MasterGo 画布业务模块多、线上问题影响设计师核心工作流的场景下，我负责团队库、DSM、变量、可交互资源、D2C 等模块的需求拆解、核心开发和线上问题闭环。",
    "针对超大文件打开和操作卡顿，我通过 Performance 定位瓶颈，并落地多级缓存、虚拟列表、细粒度组件拆分、分层分块渲染和非视口裁剪，减少无效渲染带来的性能损耗。",
    "面对 D2C 历史逻辑重复、规则配置僵硬的问题，我重构 DSL 并引入 AI 辅助生成机制，减少模板硬编码，最终使相关代码量下降约30%。",
]:
    add_bullet(doc, item)

add_role(
    doc,
    "杭州美个朋友",
    "前端负责人",
    "2023.03 - 2024.03",
    "从0到1建设多业务线前端工程体系，负责官网、后台、社媒项目、CI/CD与提效工具。"
)
for item in [
    "在公司前端体系从0到1建设阶段，我制定开发规范、CI/CD文档、项目文档和技术方案，并负责项目排期、风险跟踪与交付质量把控，保障多业务线按节奏上线。",
    "针对证书维护、异常排查、跨端兼容等重复成本高的问题，我建设证书自动替换、Web异常监控和跨端性能优化工具，降低人工维护与线上排查成本。",
]:
    add_bullet(doc, item)

add_role(
    doc,
    "杭州一知智能",
    "高级前端 / 前端小组长",
    "2021.03 - 2023.03",
    "主导多业务线前端架构改造、企业组件库和通用技术库建设。"
)
for item in [
    "面对多仓库分散、规范不统一、公共代码难复用的问题，我主导将历史项目整合为 pnpm workspace Monorepo，统一工程规范、依赖治理和共享包发布方式。",
    "在新老技术栈并存且无法一次性重构的约束下，我落地 qiankun 微前端，支持 Vue2、Vue3、React 平滑共存，并建设 yiwise-design 双版本组件库与 Storybook 文档。",
    "通过 Design Token、Atomic CSS、yiwise-utils 等公共资产治理，我推动 10+ 产品视觉和交互规范收敛，新组件开发效率提升约30%，跨项目复用成本下降约50%。",
]:
    add_bullet(doc, item)

add_role(doc, "深圳智扬信达", "前端开发工程师", "2018.10 - 2021.03", "负责基础前端项目开发、页面渲染、进度协作与代码重构。")

add_section(doc, "核心项目")
add_para(doc, "Buzzy - AI视频生成平台画布生态 | Vue3 / LeaferJS / SSE / WebSocket", size=9.1, bold=True, color="111827", after=0.6)
add_topic(
    doc,
    1,
    "Agent流式交互与画布联动",
    "AI生成链路包含文案、图片、视频和节点状态，用户如果只能等待最终结果，理解成本高且缺少过程反馈。",
    [
        "设计并实现 Agent 对话框 SSE / WebSocket 流式渲染，将模型响应、节点创建、节点更新和画布反馈串成一条链路。",
        "打通用户输入意图到 UI 工具调用的联动，使对话框可以驱动画布节点变化，而不是只返回文本结果。",
    ],
    "生成流程从黑盒等待变成节点化、可视化的过程反馈，用户能更清楚理解当前进度和下一步操作。"
)
add_topic(
    doc,
    2,
    "画布连接线与Canvas Player能力建设",
    "画布节点类型多、状态变化频繁，连接关系和媒体播放需要同时兼顾交互准确性、渲染性能和后续复用。",
    [
        "基于 LeaferJS 设计动态自适应连接线能力，处理节点移动、缩放、层级变化下的连接点计算和路径更新。",
        "设计 Canvas Player 承接图片/视频等媒体节点渲染，并将连接线、播放器能力抽象为可复用插件。",
    ],
    "沉淀 leafer-connector、leafer-player 两个插件并被 LeaferJS 官网推荐，反向验证方案具备社区复用价值。"
)

add_para(doc, "MasterGo - 协同设计软件画布性能与D2C重构 | Vue2.7 / Vue3 / Wasm / Pinia", size=9.1, bold=True, color="111827", after=0.6)
add_topic(
    doc,
    1,
    "超大文件画布性能优化",
    "超大设计文件下，团队库列表、图层播放和画布操作容易出现卡顿，影响设计师核心工作流。",
    [
        "通过 Performance 定位主耗时链路，拆分数据加载、组件渲染和画布绘制瓶颈。",
        "落地多级缓存、虚拟列表、细粒度组件拆分、分层分块渲染和非视口裁剪，减少无效渲染压力。",
    ],
    "复杂画布 FPS 稳定性明显提升，超大文件场景下的可操作性和稳定性改善。"
)
add_topic(
    doc,
    2,
    "复杂坐标系联动定位",
    "评论点、可交互资源等元素需要跟随画布缩放、拖拽和多层嵌套图层变化，坐标不一致会导致定位漂移。",
    [
        "梳理世界坐标、视图坐标、图层局部坐标和嵌套相对坐标之间的转换关系。",
        "封装坐标转换与更新链路，保障评论点、可交互资源在缩放和拖拽场景下精准联动。",
    ],
    "降低画布交互定位类问题，提升设计资产在复杂编辑场景下的稳定性。"
)
add_topic(
    doc,
    3,
    "D2C模块重构",
    "旧版 D2C 规则配置重、模板硬编码多，历史重复逻辑影响后续框架代码生成能力扩展。",
    [
        "清理旧逻辑并重构轻量 DSL，将固定模板生成改为 DSL 描述与 AI 辅助生成结合。",
        "抽离生成链路中的公共能力，降低业务规则和框架模板之间的耦合。",
    ],
    "相关代码量下降约30%，后续新增生成能力和维护历史逻辑的成本降低。"
)

add_para(doc, "super-aicc - 多技术栈融合与设计资产治理 | Monorepo / pnpm / qiankun / Vue / React", size=9.1, bold=True, color="111827", after=0.6)
add_topic(
    doc,
    1,
    "Monorepo与微前端架构升级",
    "10+ 业务项目长期独立演进，Vue/React 多技术栈共存，导致规范割裂、依赖重复和公共能力复用困难。",
    [
        "引入 pnpm workspace 聚合多业务项目，统一 lint、目录规范、依赖治理和共享包发布方式。",
        "落地 qiankun 微前端支持新老应用灰度共存，降低一次性重构风险。",
    ],
    "项目治理方式从多仓分散转为统一底座，跨项目协作和新人接入成本下降。"
)
add_topic(
    doc,
    2,
    "组件库与设计资产沉淀",
    "公共组件职责边界混乱，样式硬编码多，设计资产无法稳定复用到不同产品线。",
    [
        "建设 yiwise-design Vue/React 双版本组件库和 Storybook 文档，统一组件开发、调试和展示方式。",
        "落地 Design Token、Atomic CSS、yiwise-utils 等公共资产，收敛多产品线视觉和交互规范。",
    ],
    "新组件开发效率提升约30%，跨项目组件复用成本下降约50%。"
)

add_para(doc, "ai-code-reviewer - AI自动化代码评审工具 | Node.js / GitLab Webhook / GitLab API / LLM", size=9.1, bold=True, color="111827", after=0.6)
add_topic(
    doc,
    1,
    "AI代码评审流程落地",
    "团队 MR 数量增长后，人工 CR 容易被重复规范问题消耗，且不同成员 Review 尺度不一致。",
    [
        "设计 GitLab Webhook 监听、Event Handler、GitLab API Hook 和 Prompt 上下文组装，将差异代码与团队规范输入模型。",
        "加入评论过滤逻辑，减少无效和幻觉评论，并将有效建议自动回写到对应代码行。",
    ],
    "把质量底线前置到提交流程，降低重复 CR 成本，并形成可持续维护的开源项目。"
)

add_section(doc, "开源与文章")
for item in [
    "Leafer插件：leafer-connector、leafer-player - https://github.com/rideWind97/leafer-connector | https://github.com/rideWind97/leafer-player",
    "ai-code-reviewer - https://github.com/rideWind97/ai-code-reviewer",
    "Yiwise Design组件库：Vue版 / React版 Storybook 文档",
    "《Monorepo + Pnpm + 组件化 + CSS原子化重构实践》- https://juejin.cn/post/7254097346761162809",
]:
    add_bullet(doc, item)

doc.save(OUT)
print(OUT)
